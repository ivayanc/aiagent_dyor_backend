import requests
from typing import Dict, Any, Optional
from docx import Document
import logging
import re
import json
from agents.openai import OpenAI


logger = logging.getLogger(__name__)
logger.addFilter(lambda record: setattr(record, 'msg', f'ReportParser: {record.msg}') or True)

PARSE_DOCUMENT_EXAMPLE = '''
{  "document_title": "Mantra DYOR",  "metadata": {    "source": null,    "date_created": null,    "author": null,    "version": null  },  "summary": "Mantra is an RWA, Layer-1 focused blockchain that enables flexibility over applications to bring real-world assets on-chain while ensuring regulatory compliance.",  "sections": [    {      "section_title": "Overview",      "content": "Mantra is an RWA, Layer-1 Focused blockchain which uses a permissionless blockchain solution to enable flexibility over applications according to the Needs in the given moment ensuring & simplifying the process of bringing real-world assets on-chain but also ensures regulatory compliance for these digital representations.",      "subsections": []    },    {      "section_title": "Roadmap",      "content": "Developer Testnet, Private Testnet, Incentivized Testnet, Audit, Incubator Program, Hackathons, Mainnet Launch, DEX & CEX, Mantra Finance Integration, Bug Bounty Program, Foundation Grant, Developer Portal.",      "subsections": [        {          "subsection_title": "Hackathons",          "content": "(https://mantrachain.medium.com/mantra-announces-rwa-focused-hackathon-unleashing-innovation-in-real-world-asset-tokenization-f42e7de4e1f3)"        },        {          "subsection_title": "Mainnet Launch",          "content": "(https://www.zand.ae/en)(https://mag.ae/)(https://www.newswire.com/news/mantra-launches-mainnet-to-tokenize-real-world-assets-22446197)"        }      ]    },    {      "section_title": "Development Status",      "content": "Mantra Chain has released its mainnet, it's GitHub shows constant feedback between Users & Developers across the Ecosystem.",      "subsections": []    },    {      "section_title": "Competitors and Prospects",      "content": "Solana, Ethereum, TokenFi, LEOX, Belong.net, Brickken, Hedera.",      "subsections": []    },    {      "section_title": "Cyber Security and Audits",      "content": "Audits from SlowMist and Hacken.",      "subsections": [        {          "subsection_title": "SlowMist Audit",          "content": "(https://www.slowmist.com/security-audit-certificate.html?id=1193264f67af7c2bb840306b82eff6218471cf4fbed79a7d48d6a01a93030e35)"        },        {          "subsection_title": "Hacken Audit",          "content": "(https://audits.hacken.io/mantra-chain/)"        }      ]    },    {      "section_title": "Key Team Members",      "content": "John Patrick Mullin, Jayant Ramanand, Stephen Peepels, Nicholas Krapels, Matthew Crooks.",      "subsections": [        {          "subsection_title": "John Patrick Mullin | CEO & Founder",          "content": "LinkedIn(https://www.linkedin.com/in/john-patrick-mullin/), Twitter(https://twitter.com/jp_mullin888)"        },        {          "subsection_title": "Jayant Ramanand | Co-Founder",          "content": "LinkedIn(https://www.linkedin.com/in/jayant-ramanand/)"        },        {          "subsection_title": "Stephen Peepels | Chief Legal Officer",          "content": "LinkedIn(https://www.linkedin.com/in/stephen-peepels-ba4b2016/)"        },        {          "subsection_title": "Nicholas Krapels | CSO & Founding Advisor",          "content": "LinkedIn(https://www.linkedin.com/in/nicholaskrapels/)"        },        {          "subsection_title": "Matthew Crooks | CTO",          "content": "LinkedIn(https://www.linkedin.com/in/mcrooks/)"        }      ]    },    {      "section_title": "Reputation",      "content": "MANTRA has established a notable presence in the decentralized finance (DeFi) sector, particularly in tokenizing real-world assets (RWAs).",      "subsections": []    },    {      "section_title": "Social Media Presence",      "content": "Medium, Telegram, Youtube, Discord, Instagram, Twitter.",      "subsections": [        {          "subsection_title": "Medium",          "content": "(https://mantrachain.medium.com/)"        },        {          "subsection_title": "Telegram",          "content": "(https://t.me/MANTRA_Chain)"        },        {          "subsection_title": "Youtube",          "content": "(https://www.youtube.com/@MANTRAChain)"        },        {          "subsection_title": "Discord",          "content": "(https://discord.com/invite/mantrachain)"        },        {          "subsection_title": "Instagram",          "content": "(https://www.instagram.com/mantra_chain/)"        },        {          "subsection_title": "Twitter",          "content": "(https://twitter.com/MANTRA_Chain)"        }      ]    },    {      "section_title": "Business Model",      "content": "MANTRA's revenue is generated through transaction fees, DeFi services fees, and partnerships for asset tokenization.",      "subsections": []    },    {      "section_title": "Investments",      "content": "Involves various funding rounds with undisclosed amounts raised.",      "subsections": [        {          "subsection_title": "IDA Funding Round",          "content": "Date: 15 Aug 2020, Raised: ?, Investors: Genblock, Kenetic Capital, LD Capital, Waterdrip Capital, Moonrock Capital, Master Ventures, CSP DAO, Vendetta Capital, Caladan, Fairum Ventures, Plutus VC, DeltaHub Capital, Sky Ventures."        },        {          "subsection_title": "Undisclosed Round",          "content": "Date: 19 Mar 2024, Raised: $ 11.00M, Investors: Mapeblock Capital, Token Bay Capital, GameFi Ventures, Forte, BlackPine, Hex Trust, 280 Capital, Fuse Capital, Three Point Capital, Virtuzone, Shorooq Partners."        },        {          "subsection_title": "Strategic Round",          "content": "Date: 23 May 2024, Raised: ?, Investors: Laser Digital."        }      ]    },    {      "section_title": "Conclusion",      "content": "Mantra shows as a Solution for RWA Industries who want to introduce into the Tokenization of Real State.",      "subsections": []    }  ],  "team": [    {      "name": "John Patrick Mullin",      "role": "CEO & Founder",      "bio": "An entrepreneur focused on blockchain technology with over 10 years of experience.",      "contacts": {        "linkedin": "https://www.linkedin.com/in/john-patrick-mullin/",        "twitter": "https://twitter.com/jp_mullin888",        "other": null      }    },    {      "name": "Jayant Ramanand",      "role": "Co-Founder",      "bio": "Graduated in Economics with an MBA, operational & infrastructure lead at Mantra.",      "contacts": {        "linkedin": "https://www.linkedin.com/in/jayant-ramanand/",        "twitter": null,        "other": null      }    },    {      "name": "Stephen Peepels",      "role": "Chief Legal Officer",      "bio": "Has over 15 years of experience advising on complex international transactions.",      "contacts": {        "linkedin": "https://www.linkedin.com/in/stephen-peepels-ba4b2016/",        "twitter": null,        "other": null      }    },    {      "name": "Nicholas Krapels",      "role": "CSO & Founding Advisor",      "bio": "An American writer and entrepreneur with extensive experience in China.",      "contacts": {        "linkedin": "https://www.linkedin.com/in/nicholaskrapels/",        "twitter": null,        "other": null      }    },    {      "name": "Matthew Crooks",      "role": "CTO",      "bio": "With over 25 years of experience in insurance and financial services, he excels in bridging technology and business.",      "contacts": {        "linkedin": "https://www.linkedin.com/in/mcrooks/",        "twitter": null,        "other": null      }    }  ],  "social_media": {    "platforms": [      {        "name": "Medium",        "url": "https://mantrachain.medium.com/",        "followers": null      },      {        "name": "Telegram",        "url": "https://t.me/MANTRA_Chain",        "followers": "129.7k Members"      },      {        "name": "Youtube",        "url": "https://www.youtube.com/@MANTRAChain",        "followers": "14.9k Subscribers"      },      {        "name": "Discord",        "url": "https://discord.com/invite/mantrachain",        "followers": "165k Members"      },      {        "name": "Instagram",        "url": "https://www.instagram.com/mantra_chain/",        "followers": null      },      {        "name": "Twitter",        "url": "https://twitter.com/MANTRA_Chain",        "followers": "321k Followers"      }    ]  },  "investments": [    {      "round_name": "IDA Funding Round",      "date": "15 Aug 2020",      "amount_raised": "?",      "investors": [        {          "name": "Genblock Capital",          "url": "https://cryptorank.io/funds/genblock-capital/rounds"        },        {          "name": "Kenetic Capital",          "url": "https://cryptorank.io/funds/kenetic-capital/rounds"        },        {          "name": "LD Capital",          "url": "https://cryptorank.io/funds/ld-capital/rounds"        },        {          "name": "Waterdrip Capital",          "url": "https://cryptorank.io/funds/waterdrip-capital/rounds"        },        {          "name": "Moonrock Capital",          "url": "https://cryptorank.io/funds/moonrock-capital/rounds"        },        {          "name": "Master Ventures",          "url": "https://cryptorank.io/funds/master-ventures/rounds"        },        {          "name": "CSP DAO",          "url": "https://cryptorank.io/funds/csp-dao/rounds"        },        {          "name": "Vendetta Capital",          "url": "https://cryptorank.io/funds/vendetta-capital/rounds"        },        {          "name": "Caladan",          "url": "https://cryptorank.io/funds/alphalab-capital/rounds"        },        {          "name": "Fairum Ventures",          "url": "https://cryptorank.io/funds/fairum-ventures/rounds"        },        {          "name": "Plutus VC",          "url": "https://cryptorank.io/funds/plutus-vc/rounds"        },        {          "name": "DeltaHub Capital",          "url": "https://cryptorank.io/funds/deltahub-capital/rounds"        },        {          "name": "Sky Ventures",          "url": "https://cryptorank.io/funds/sky-ventures/rounds"        }      ]    },    {      "round_name": "Undisclosed Round",      "date": "19 Mar 2024",      "amount_raised": "$ 11.00M",      "investors": [        {          "name": "Mapeblock Capital",          "url": "https://cryptorank.io/funds/mapleblock-capital/rounds"        },        {          "name": "Token Bay Capital",          "url": "https://cryptorank.io/funds/tokenbaycapital/rounds"        },        {          "name": "GameFi Ventures",          "url": "https://cryptorank.io/funds/gamefi-ventures/rounds"        },        {          "name": "Forte",          "url": "https://cryptorank.io/funds/forte/rounds"        },        {          "name": "BlackPine",          "url": "https://cryptorank.io/funds/blackpine/rounds"        },        {          "name": "Hex Trust",          "url": "https://cryptorank.io/funds/hex-trust/rounds"        },        {          "name": "280 Capital",          "url": "https://cryptorank.io/funds/280-capital/rounds"        },        {          "name": "Fuse Capital",          "url": "https://cryptorank.io/funds/fuse-capital/rounds"        },        {          "name": "Three Point Capital",          "url": "https://cryptorank.io/funds/three-point/rounds"        },        {          "name": "Virtuzone",          "url": "https://cryptorank.io/funds/virtuzone/rounds"        },        {          "name": "Shorooq Partners",          "url": "https://cryptorank.io/funds/shorooq-partners/rounds"        }      ]    },    {      "round_name": "Strategic Round",      "date": "23 May 2024",      "amount_raised": "?",      "investors": [        {          "name": "Laser Digital",          "url": "https://cryptorank.io/funds/laser-digital/rounds"        }      ]    }  ],  "additional_info": {    "notes": "There are some flaws that should be covered to attract investors.",    "references": []  }}
'''

DYOR_JSON_STRUCTURE = {
  "document_title": "",
  "metadata": {
    "source": "",
    "date_created": "",
    "author": "",
    "version": "",
  },
  "general_info": {
    "project_name": "",
    "research_date": "",
    "token_info": {
      "token_chain": "",
      "token_address": "",
    },
    "github_url": "",
  },
  "summary": "",
  "sections": [
    {
      "section_title": "",
      "content": "",
      "subsections": [
        {
          "subsection_title": "",
          "content": ""
        }
      ]
    }
  ],
  "team": [
    {
      "name": "",
      "role": "",
      "bio": "",
      "contacts": {
        "linkedin": "",
        "twitter": "",
        "other": ""
      }
    }
  ],
  "social_media": {
    "platforms": [
      {
        "name": "",
        "url": "",
        "followers": ""
      }
    ]
  },
  "investments": [
    {
      "round_name": "",
      "date": "",
      "amount_raised": "",
      "investors": [
        {
          "name": "",
          "url": ""
        }
      ]
    }
  ],
  "additional_info": {
    "notes": "",
    "references": [
      {
        "description": "",
        "url": ""
      }
    ]
  }
}


class DYORParser(OpenAI):
    DEFAULT_LORE = f"""You are a DYOR (Do Your Own Research) report parser.
                    Your task is to extract structured information from cryptocurrency project research reports.
                    Parse the given text and return only the requested JSON structure with relevant URLs and data.
                    If a field is not found, use null instead of leaving it empty.
                    
                    Use the following JSON structure:
                    {json.dumps(DYOR_JSON_STRUCTURE, indent=1)}"""

    def __init__(self, api_key: str):
        self.api_key = api_key
        self.base_url = "https://api.openai.com/v1"
        self.headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }

    def parse_document(self, file_path: str) -> str:
        """
        Parse a DOCX document into text with hyperlinks formatted as HyperText.
        Links are only processed within their original paragraphs.
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            str: Document text with hyperlinks formatted as text(url)
        """
        doc = Document(file_path)
        processed_paragraphs = []
        logger.error(f"Processing {len(doc.paragraphs)} paragraphs")
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if len(paragraph.hyperlinks) > 0:
              logger.error(f"Found {len(paragraph.hyperlinks)} hyperlinks")
              for hyperlink in paragraph.hyperlinks:
                  link_text = hyperlink.text
                  url = hyperlink.url
                  if link_text in text:
                      start = text.find(link_text)
                      if start >= 0:
                          before = text[:start]
                          after = text[start + len(link_text):]
                          text = before + f"{link_text}({url})" + after
                  else:
                      text = text + f"{link_text}({url})"
              processed_paragraphs.append(text)
                  
        return "\n".join(processed_paragraphs)

    def parse_document_with_openai(self, file_path: str) -> str:
        from pprint import pprint
        parsed_text = self.parse_document(file_path)
        prompt = (f"Parse the following DOCX document and return only the requested JSON structure with relevant URLs and data. "
                  f"If a field is not found, use null instead of leaving it empty.\n\n{parsed_text}")
        res = self.chat(prompt)
        
        pprint(res)
        json_res = self.parse_json(res)
        return json_res

    def parse_json(self, json_text = None):
        if not json_text:
            json_text = PARSE_DOCUMENT_EXAMPLE
        return json.loads(json_text)